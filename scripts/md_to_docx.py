#!/usr/bin/env python3
"""
Learn Everything — Markdown → DOCX 批量转换脚本

用法：
  python3 scripts/md_to_docx.py                  # 转换所有模块
  python3 scripts/md_to_docx.py --module 09      # 只转换指定模块
  python3 scripts/md_to_docx.py --domain ai-product-manager  # 指定领域

图片嵌入：
  读取 domains/{slug}/assets/manifest.md，按 manifest 将图片插入对应段落之后。

依赖：
  pip install python-docx pillow
"""

import argparse
import os
import re
import csv
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image as PILImage
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

DOMAINS_DIR = Path(__file__).parent.parent / 'domains'

# 模块编号 → DOCX 输出文件名
FILENAME_MAP = {
    '00': '00-学习路径（先读这个）',
    '01': '01-产品管理基础',
    '02': '02-AI与ML技术素养',
    '03': '03-人机交互设计（HCAI）',
    '04': '04-LLM与GenAI产品管理',
    '05': '05-Agentic AI产品设计',
    '06': '06-AI产品评估体系',
    '07': '07-AI产品战略与GTM',
    '08': '08-AI成本经济学',
    '09': '09-MLOps与运维现实',
    '10': '10-AI伦理与合规',
    '11': '11-教育AI的特殊考量',
    '12': '12-AI PM的持续演进',
    '13': '13-Agentic产品案例分析',
    '14': '14-Agent推理与规划框架',
    '15': '15-Agent记忆系统与工具调用',
}


# ──────────────────────────────────────────────
# 样式工具函数
# ──────────────────────────────────────────────

def set_page_layout(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def add_run_with_inline(para, text, base_size=10.5, base_bold=False):
    """解析行内 **bold** 和 `code`，添加带格式的 run。"""
    pattern = r'(\*\*.*?\*\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.font.name = 'Calibri'
            run.font.size = Pt(base_size)
            run.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(31, 73, 125)
        else:
            run = para.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(base_size)
            run.font.bold = base_bold


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 73, 125)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def flush_table(doc, table_data):
    if not table_data:
        return
    rows = [r for r in table_data
            if not all(re.match(r'^[-:\s]+$', c) for c in r if c)]
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = 'Table Grid'
    for ri, row_cells in enumerate(rows):
        for ci in range(col_count):
            cell_text = row_cells[ci] if ci < len(row_cells) else ''
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            add_run_with_inline(p, cell_text.strip(),
                                base_size=9.5, base_bold=(ri == 0))
    doc.add_paragraph()


# ──────────────────────────────────────────────
# 核心解析器
# ──────────────────────────────────────────────

def parse_and_build(md_text: str, doc: Document) -> list[tuple[str, object]]:
    """
    解析 Markdown 并构建 DOCX。
    返回 [(paragraph_text, paragraph_object), ...] 供图片嵌入使用。
    """
    set_page_layout(doc)
    lines = md_text.split('\n')
    para_index: list[tuple[str, object]] = []

    # 跳过 YAML frontmatter
    i = 0
    if lines and lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1

    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_data: list[list[str]] = []

    def do_flush_table():
        nonlocal in_table, table_data
        flush_table(doc, table_data)
        in_table = False
        table_data = []

    while i < len(lines):
        line = lines[i]

        # ── 代码块 ──
        if line.strip().startswith('```'):
            if in_code_block:
                p = add_code_block(doc, '\n'.join(code_lines))
                para_index.append(('\n'.join(code_lines), p))
                in_code_block = False
                code_lines = []
            else:
                if in_table:
                    do_flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── 表格行 ──
        if line.strip().startswith('|') and line.strip().endswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(re.match(r'^[-:\s]+$', c) for c in cells if c):
                i += 1
                continue
            in_table = True
            table_data.append(cells)
            i += 1
            continue
        else:
            if in_table:
                do_flush_table()

        # ── 空行 ──
        if not line.strip():
            i += 1
            continue

        # ── 标题 ──
        if line.startswith('#### '):
            p = doc.add_heading(level=4)
            p.clear()
            add_run_with_inline(p, line[5:].strip(), base_size=10.5, base_bold=True)
            para_index.append((line[5:].strip(), p))
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            p.clear()
            add_run_with_inline(p, line[4:].strip(), base_size=12, base_bold=True)
            para_index.append((line[4:].strip(), p))
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            p.clear()
            add_run_with_inline(p, line[3:].strip(), base_size=14, base_bold=True)
            para_index.append((line[3:].strip(), p))
        elif line.startswith('# '):
            p = doc.add_heading(level=1)
            p.clear()
            add_run_with_inline(p, line[2:].strip(), base_size=18, base_bold=True)
            para_index.append((line[2:].strip(), p))
        # ── 水平线 ──
        elif re.match(r'^-{3,}$', line.strip()):
            pass
        # ── 列表：缩进子弹 ──
        elif re.match(r'^\s{2,}[-*] ', line):
            p = doc.add_paragraph(style='List Bullet 2')
            text = re.sub(r'^\s+[-*] ', '', line)
            add_run_with_inline(p, text.strip(), base_size=10)
            para_index.append((text.strip(), p))
        # ── 列表：一级子弹 ──
        elif re.match(r'^[-*] ', line):
            p = doc.add_paragraph(style='List Bullet')
            add_run_with_inline(p, line[2:].strip(), base_size=10.5)
            para_index.append((line[2:].strip(), p))
        # ── 列表：有序 ──
        elif re.match(r'^\d+[\.\)] ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+[\.\)] ', '', line)
            add_run_with_inline(p, text.strip(), base_size=10.5)
            para_index.append((text.strip(), p))
        # ── 普通段落 ──
        else:
            p = doc.add_paragraph(style='Normal')
            add_run_with_inline(p, line.strip(), base_size=10.5)
            para_index.append((line.strip(), p))

        i += 1

    if in_table:
        do_flush_table()

    return para_index


# ──────────────────────────────────────────────
# 图片嵌入
# ──────────────────────────────────────────────

def load_manifest(assets_dir: Path) -> list[dict]:
    manifest_path = assets_dir / 'manifest.md'
    if not manifest_path.exists():
        return []
    entries = []
    with open(manifest_path, encoding='utf-8') as f:
        in_table = False
        headers = []
        for line in f:
            line = line.strip()
            if line.startswith('|') and line.endswith('|'):
                cells = [c.strip().strip('`') for c in line.strip('|').split('|')]
                if all(re.match(r'^[-:\s]+$', c) for c in cells if c):
                    continue
                if not headers:
                    headers = [c.lower().replace(' ', '_') for c in cells]
                    in_table = True
                elif in_table:
                    row = dict(zip(headers, cells))
                    entries.append(row)
    return entries


def embed_images(doc: Document, para_index: list[tuple[str, object]],
                 assets_dir: Path, module_stem: str):
    """根据 manifest 将图片嵌入到对应段落后面。"""
    manifest = load_manifest(assets_dir)
    if not manifest:
        return

    for entry in manifest:
        filename = entry.get('文件名', '')
        belongs_to = entry.get('归属模块', '')
        keyword = entry.get('插入位置（段落关键词）', '')
        try:
            width_cm = float(entry.get('宽度(cm)', '13.0'))
        except ValueError:
            width_cm = 13.0

        # 检查图片是否属于当前模块
        if belongs_to and module_stem not in belongs_to and belongs_to not in module_stem:
            continue

        img_path = assets_dir / filename
        if not img_path.exists():
            print(f'  [WARN] 图片不存在: {img_path}')
            continue

        # 找到关键词所在的段落
        target_para = None
        for text, p in para_index:
            if keyword and keyword in text:
                target_para = p
                break

        if target_para is None:
            # 找不到关键词，插在文档末尾
            print(f'  [WARN] 未找到插入关键词"{keyword}"，图片 {filename} 插入文档末尾')

        # 在目标段落后插入图片段落
        try:
            # 创建图片段落
            img_para = OxmlElement('w:p')
            img_para_pr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            img_para_pr.append(jc)
            img_para.append(img_para_pr)

            # 使用 python-docx 添加图片（简单方式）
            p = doc.add_paragraph()
            p.alignment = 1  # center
            run = p.add_run()
            run.add_picture(str(img_path), width=Cm(width_cm))

            # 将段落移到目标位置
            if target_para is not None:
                img_elem = p._element
                doc.element.body.remove(img_elem)
                target_para._element.addnext(img_elem)
            # else: 留在末尾

            print(f'  [IMG] 嵌入 {filename} → "{keyword}"')
        except Exception as e:
            print(f'  [ERROR] 嵌入图片失败 {filename}: {e}')


# ──────────────────────────────────────────────
# 主流程
# ──────────────────────────────────────────────

def convert_domain(domain_slug: str, only_module: str = None):
    domain_dir = DOMAINS_DIR / domain_slug
    modules_dir = domain_dir / 'modules'
    output_dir = domain_dir / 'published'
    assets_dir = domain_dir / 'assets'

    output_dir.mkdir(exist_ok=True)

    md_files = sorted(modules_dir.glob('*.md'))
    if not md_files:
        print(f'[WARN] {domain_slug}: modules/ 目录为空')
        return

    ok_count = 0
    for md_path in md_files:
        stem = md_path.stem
        num = stem.split('-')[0]

        if only_module and num != only_module.zfill(2):
            continue

        if num not in FILENAME_MAP:
            print(f'SKIP  {md_path.name} (编号 {num} 不在 FILENAME_MAP)')
            continue

        out_name = FILENAME_MAP[num] + '.docx'
        out_path = output_dir / out_name

        print(f'转换  {md_path.name} → {out_name}')
        md_text = md_path.read_text(encoding='utf-8')
        doc = Document()
        para_index = parse_and_build(md_text, doc)

        # 嵌入图片
        if assets_dir.exists():
            embed_images(doc, para_index, assets_dir, stem)

        doc.save(out_path)
        ok_count += 1

    print(f'\n完成：{ok_count} 个文件 → {output_dir}')


def main():
    parser = argparse.ArgumentParser(description='Markdown → DOCX 批量转换')
    parser.add_argument('--domain', default='ai-product-manager',
                        help='领域 slug（默认：ai-product-manager）')
    parser.add_argument('--module', default=None,
                        help='只转换指定模块编号，如 --module 09')
    args = parser.parse_args()

    convert_domain(args.domain, args.module)


if __name__ == '__main__':
    main()
